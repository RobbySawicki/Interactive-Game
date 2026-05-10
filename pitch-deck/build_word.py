"""
Build the editable Word version of the Interactive Media pitch deck.

Each "slide" is one landscape page so the doc reads the same as the PDF.
Designed for clients to NOT see this directly — this is the editable copy
the sales rep / brand owner updates between pitches.

Run:  python3 build_word.py
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ──────────────────────────────────────────────────────────────────────────────
# Palette (matches the PDF / site)
# ──────────────────────────────────────────────────────────────────────────────
BG      = "0b0b0a"
BG2     = "141312"
INK     = "f4f1ea"
INK2    = "b8b3a7"
INK3    = "6a6557"
LINE    = "2a2825"
ACCENT  = "d4e34a"


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_page_bg(doc, color_hex):
    """Add a Word page background color (visible in 'Web Layout' but doesn't
    print — good enough for an editable working file)."""
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), color_hex)
    doc.element.insert(0, bg)
    display = OxmlElement("w:displayBackgroundShape")
    settings = doc.settings.element
    settings.append(display)


def add_run(p, text, *, font="Helvetica", size=11, bold=False, italic=False,
            color=INK, all_caps=False, letter_spacing=None):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = RGBColor.from_string(color)
    if all_caps:
        r.font.all_caps = True
    # letter spacing via xml
    if letter_spacing is not None:
        rPr = r._element.get_or_add_rPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:val"), str(letter_spacing))
        rPr.append(spacing)
    return r


def landscape_section(doc):
    sec = doc.sections[-1]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Inches(13.333)
    sec.page_height = Inches(7.5)
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ──────────────────────────────────────────────────────────────────────────────
# Slide builders
# ──────────────────────────────────────────────────────────────────────────────
def slide_header(doc, slide_num, total, section_label):
    """Three-column header table: brand · section · counter."""
    tbl = doc.add_table(rows=1, cols=3)
    tbl.autofit = False
    tbl.columns[0].width = Inches(3.5)
    tbl.columns[1].width = Inches(5.5)
    tbl.columns[2].width = Inches(2.5)

    # Brand cell
    bcell = tbl.cell(0, 0)
    set_cell_bg(bcell, BG)
    p = bcell.paragraphs[0]
    add_run(p, "◼ ", font="Helvetica", size=10, bold=True, color=ACCENT)
    add_run(p, "INTERACTIVE MEDIA", font="Helvetica", size=9, bold=True,
            color=INK, all_caps=True, letter_spacing=40)

    # Section label
    scell = tbl.cell(0, 1)
    set_cell_bg(scell, BG)
    sp = scell.paragraphs[0]
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sp, section_label.upper() if section_label else "",
            font="Courier New", size=8, color=INK3, letter_spacing=40)

    # Slide counter
    ccell = tbl.cell(0, 2)
    set_cell_bg(ccell, BG)
    cp = ccell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(cp, f"{slide_num:02d} / {total:02d}",
            font="Courier New", size=8, color=INK3, letter_spacing=40)

    # accent line under header
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(2)
    rule.paragraph_format.space_after = Pt(8)
    pPr = rule._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), LINE)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_eyebrow(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "▬ ", font="Helvetica", size=10, color=ACCENT, bold=True)
    add_run(p, text.upper(), font="Courier New", size=9, color=INK2,
            letter_spacing=60)


def add_title(doc, text, size=44):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    # support {em}…{em} for italic INK2
    parts = text.split("{em}")
    for i, chunk in enumerate(parts):
        if i % 2 == 0:
            add_run(p, chunk, font="Helvetica", size=size, bold=True, color=INK)
        else:
            add_run(p, chunk, font="Garamond", size=size, italic=True, color=INK2)


def add_lede(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    add_run(p, text, font="Helvetica", size=14, color=INK2)


def add_section_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text.upper(), font="Courier New", size=9, bold=True,
            color=ACCENT, letter_spacing=60)


def add_bullet(doc, kw, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p, kw.upper() + "   ", font="Courier New", size=9, bold=True,
            color=ACCENT, letter_spacing=40)
    add_run(p, body, font="Helvetica", size=11, color=INK2)


def add_card_row(doc, cards, cols=None):
    """cards = [(kw, headline, body), ...] — rendered as a colored row of cells."""
    if cols is None:
        cols = len(cards)
    rows_needed = (len(cards) + cols - 1) // cols
    tbl = doc.add_table(rows=rows_needed, cols=cols)
    tbl.autofit = False
    for col in range(cols):
        tbl.columns[col].width = Inches(11.6 / cols)
    for i, (kw, head, body) in enumerate(cards):
        r = i // cols
        col = i % cols
        cell = tbl.cell(r, col)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_bg(cell, BG2)
        # clear default empty paragraph spacing
        cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
        # accent label
        p1 = cell.add_paragraph()
        p1.paragraph_format.space_before = Pt(6)
        p1.paragraph_format.space_after = Pt(2)
        add_run(p1, kw.upper(), font="Courier New", size=9, bold=True,
                color=ACCENT, letter_spacing=60)
        # headline
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        add_run(p2, head, font="Helvetica", size=13, bold=True, color=INK)
        # body
        p3 = cell.add_paragraph()
        p3.paragraph_format.space_after = Pt(6)
        add_run(p3, body, font="Helvetica", size=10, color=INK2)


def add_spacer(doc, pt=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)


# ──────────────────────────────────────────────────────────────────────────────
# Slide definitions
# ──────────────────────────────────────────────────────────────────────────────
def slide_cover(doc):
    add_eyebrow(doc, "Pitch · v1.0 · 2026")
    add_title(doc, "Guerrilla OOH that {em}plays back{em}.", size=46)
    add_lede(doc,
             "A two-screen interactive platform for in-person brand activations. "
             "One player on an iPad. A crowd on a giant LED panel or projection. "
             "Live sync. Branded per campaign.")
    add_section_label(doc, "What's in this deck")
    add_bullet(doc, "01–04", "Problem, product, setup.")
    add_bullet(doc, "05–07", "The three surfaces — iPad, side LED, capture panel.")
    add_bullet(doc, "08–10", "Mechanic ideas and brand reskin examples.")
    add_bullet(doc, "11–13", "Contexts, deliverables, demo codes.")
    add_bullet(doc, "14",     "Call to action — your contact details go here.")


def slide_problem(doc):
    add_eyebrow(doc, "01 · The problem")
    add_title(doc, "OOH is loud. {em}Almost no one is listening.{em}", size=36)
    add_lede(doc,
             "Static boards get a half-second of attention. Digital boards get "
             "tuned out the same way social ads do. Brands keep paying for "
             "impressions that don't actually impress.")
    add_card_row(doc, [
        ("Glance time",   "0.4s",   "Average attention paid to a billboard."),
        ("Personalisation gap", "70%", "Consumers say outdoor ads don't feel relevant to them."),
        ("Walk-by conversion",  "9/10","Activations fail to convert without an active hook."),
    ])
    add_spacer(doc, 6)
    p = doc.add_paragraph()
    add_run(p, "Attention now has to be earned in seconds — not bought in CPMs.",
            font="Helvetica", size=14, bold=True, color=INK)


def slide_idea(doc):
    add_eyebrow(doc, "02 · The product")
    add_title(doc, "It's an {em}interactive game{em}.", size=48)
    add_lede(doc,
             "A player picks up an iPad and plays. The crowd watches the action "
             "blow up on a giant LED panel or projection beside them. Every tap, "
             "hit, miss, and score syncs across both screens live. Spectators "
             "pull into a crowd. The crowd pulls in more players.")
    add_card_row(doc, [
        ("01 · PLAYER",  "iPad in hand",   "The touch surface where the round happens. 4:3 landscape, branded skin, on-site tunable."),
        ("02 · HYPE",    "Side LED / projection", "The big screen beside the player. Score, combo, reaction words, leaderboard rail."),
        ("03 · CAPTURE", "Rear panel + QR",       "Square surface for the brand drop. Scan to enter, sample, follow, sign up."),
    ])


def slide_why(doc):
    add_eyebrow(doc, "03 · Why it works")
    add_title(doc, "Stop them. Hold them. {em}Capture them.{em}", size=36)
    add_card_row(doc, [
        ("STOP",
         "Pulls a crowd in seconds.",
         "A live game on a giant LED panel is loud, bright, and visibly alive. People stop to watch — then they stop to play."),
        ("HOLD",
         "30-second rounds, 90-second dwell.",
         "Players stay for the round. Spectators stay for the leaderboard, the next player, and the chance to take a shot themselves."),
        ("CAPTURE",
         "Turns attention into action.",
         "A QR on the rear panel or projection lets you collect emails, hand out samples, drop coupons, or launch a follow-up flow."),
    ])
    add_spacer(doc, 6)
    p = doc.add_paragraph()
    add_run(p, "ENTIRE FUNNEL — FROM \"PASSER-BY\" TO \"OPT-IN\" — ON ONE STREET CORNER.",
            font="Courier New", size=9, color=INK3, letter_spacing=60)


def slide_setup(doc):
    add_eyebrow(doc, "04 · The setup")
    add_title(doc, "Three surfaces. {em}One game.{em}", size=40)
    add_lede(doc,
             "Every event, hit, score, and miss is broadcast live across all "
             "three surfaces. The truck can be a real LED truck, a stage-side "
             "LED wall, a window display, or a wall projection. The square panel "
             "can be replaced by a kiosk, a poster, or a step-and-repeat with a "
             "QR. Same engine, different real-estate.")
    add_card_row(doc, [
        ("SIDE LED · 1007 × 432",  "Hype reel",   "Ultra-wide format. The screen that pulls the crowd in. Live score, combo, reactions, top scores."),
        ("REAR LED · 1:1",         "Capture surface", "Square format. QR drop, comic-burst pulse on every hit. Connects spectator to brand."),
        ("iPad · 4:3 touch",       "Player surface",  "Where the round happens. Touch-first mechanic, branded skin, leaderboard handoff at the end."),
    ])


def slide_ipad(doc):
    add_eyebrow(doc, "05 · Surface — the player")
    add_title(doc, "The {em}iPad{em}.", size=48)
    add_bullet(doc, "Touch-first",   "Drag, fling, tap, tilt, draw — the mechanic is built around what hands do best on a tablet.")
    add_bullet(doc, "Branded skin",  "Same engine, fully reskinned per campaign — palette, fonts, hero art, sound, copy.")
    add_bullet(doc, "Live score",    "Round timer, combo meter, reaction words — every hit feels rewarding.")
    add_bullet(doc, "End screen",    "Name entry feeds the live leaderboard on the LED. Capture happens here too.")
    add_bullet(doc, "On-site dial",  "Hidden tweaks panel lets the operator tune difficulty in real time as the crowd shifts.")
    add_bullet(doc, "Hardware",      "Runs in any iPad browser, no app, no install. Bring your own iPad or rent on-site.")


def slide_side_led(doc):
    add_eyebrow(doc, "06 · Surface — the hype")
    add_title(doc, "The {em}LED panel{em} (or projection).", size=36)
    add_bullet(doc, "Live score",       "Score updates instantly on every hit — the screen is alive in a way static OOH never is.")
    add_bullet(doc, "Reaction words",   "POW, KAPOW, NICE — timed bursts that make the screen feel like the crowd is cheering.")
    add_bullet(doc, "Combo meter",      "Visual streak indicator that escalates with the player's run. Watch one fill, you stay.")
    add_bullet(doc, "Leaderboard rail", "Top scores from the day. People come back to take down the leader. Built-in repeat play.")
    add_bullet(doc, "Deploys as",       "LED truck side, stage-side LED wall, retail window projection, building-side projection, arena ribbon.")


def slide_rear(doc):
    add_eyebrow(doc, "07 · Surface — the capture")
    add_title(doc, "The {em}rear panel{em} closes the loop.", size=32)
    add_bullet(doc, "Square format",     "1:1 panel — fits the back of a truck, a stage scrim, a window, a freestanding kiosk.")
    add_bullet(doc, "QR drop",           "Scan-to-anywhere. Email list, sample request, sweepstakes, coupon, IG follow, sign-up.")
    add_bullet(doc, "Live bursts",       "Comic-style POW bursts pulse on every hit — the panel keeps moving even between players.")
    add_bullet(doc, "Foot traffic ROI",  "This is the panel that converts the spectator into a pixel you can re-market to.")


def slide_game_ideas_action(doc):
    add_eyebrow(doc, "08 · What the game can be")
    add_title(doc, "The mechanic is {em}yours{em}.", size=40)
    add_lede(doc, "Same platform. Pick the action that fits the brand. "
                  "Examples — not a fixed menu.")
    add_card_row(doc, [
        ("SLING THE THING", "Drag-and-fling slingshot",
         "Branded projectile: a cookie, a lipstick wand, a basketball, a beer can, a bottle, a snack pack. 30s sprint, combos, scoring."),
        ("TAP TO THE BEAT", "Reflex sprint",
         "Glowing branded dots appear, tap before they fade. Combo multipliers. Music-synced. Great for energy, music, fitness."),
        ("CATCH & DODGE",   "Branded rain",
         "Catch the good (branded products), dodge the bad. Tilt or tap to move the bucket. Built-in good-vs-evil narrative."),
        ("RACE / DRIVE",    "Steer & dodge",
         "Tilt the iPad or drag to steer a branded vehicle through obstacles. Distance-based scoring."),
    ], cols=2)


def slide_game_ideas_brand(doc):
    add_eyebrow(doc, "08 · What the game can be (cont.)")
    add_title(doc, "Or skip the action. {em}Quiz them. Vote them. Style them.{em}", size=28)
    add_card_row(doc, [
        ("TRIVIA / QUIZ",   "Multiple-choice rounds",
         "Brand knowledge, sports trivia, pop-culture rounds. Live leaderboard. Sports partners, retail, media brands."),
        ("VOTE / TASTE TEST", "Head-to-head picks",
         "Original vs. Crunchy. Red vs. Blue. Crowd-sourced votes drive a live tally on the LED — spectators argue with the leaderboard."),
        ("SPIN TO WIN",     "Wheel mechanic",
         "Classic prize-wheel reskinned: sample, coupon, free product, swag. QR confirms the win on the rear panel."),
        ("DRAW / SIGN / STYLE", "Creative mode",
         "Sign a board, customize a jersey, color a product, paint a virtual graffiti wall. Output saves and emails."),
        ("PHOTO BOOTH",     "AR-style frame",
         "Camera mode with branded overlay — the photo posts to the LED panel, then gets emailed via QR scan."),
        ("UNLOCK / COMBO",  "Branded codes",
         "Collect items from multiple posters around the venue to unlock a final scene on the truck. Scavenger-hunt flavor."),
    ], cols=3)


def slide_example_walkthrough(doc):
    add_eyebrow(doc, "09 · Worked example")
    add_title(doc, "Walking through {em}one round{em}.", size=36)
    add_card_row(doc, [
        ("01", "Player taps the iPad.",
         "Name entry, then the game starts. The LED swaps from attract loop to LIVE — the crowd notices."),
        ("02", "Each hit broadcasts.",
         "Slingshot example: drag back, release. The iPad shows the throw, the LED shows the score and a POW! burst."),
        ("03", "Combos build hype.",
         "Streaks light the combo meter and trigger bigger bursts on the side LED. People stop walking."),
        ("04", "30s. Final score.",
         "Name lands on the leaderboard. LED tells the crowd if it's a new top score. Player wants another go."),
        ("05", "Spectator scans the QR.",
         "Rear panel pulls them off the sidewalk and onto your sign-up / sample / coupon page."),
    ], cols=5)


def slide_examples_brands(doc):
    add_eyebrow(doc, "10 · Reskin examples")
    add_title(doc, "Same engine. {em}Different brand, every time.{em}", size=32)
    add_card_row(doc, [
        ("SNACK BRAND",   "Sling the cookie / chip / candy bar",
         "Branded projectile and target. End screen offers a coupon or sample."),
        ("BEAUTY BRAND",  "Apply the mascara wand / lipstick",
         "Drag-and-paint mechanic. Final look exports to email."),
        ("BEER / BEVERAGE","Toss the can into the cooler",
         "Slingshot reskinned as a branded can. Crowd cheers a perfect landing."),
        ("SPORTS TEAM",   "Free throws / penalty kicks",
         "Branded ball, team court, mascot reactions. Leaderboard sponsored by partner."),
        ("AUTO BRAND",    "Drive the new model",
         "Tilt-to-steer. Truck shows your run. Leaderboards the fastest lap of the day."),
        ("STREAMER / MEDIA","Trivia from a show / franchise",
         "Multiple-choice rounds, instant scoring, sweepstakes drop on the rear QR."),
        ("FAST FOOD",     "Build-the-burger / fry-toss",
         "Drag ingredients into the bun in 30 seconds. Closest order wins a free meal."),
        ("RETAILER",      "Catch the falling items into the cart",
         "Catch & dodge reskin. Spend tally on the LED. Coupon at end."),
    ], cols=4)


def slide_contexts(doc):
    add_eyebrow(doc, "11 · Where it goes")
    add_title(doc, "Anywhere people {em}gather{em}.", size=44)
    add_card_row(doc, [
        ("FESTIVALS",        "Concerts, food fests, fairs.",  "Crowds with downtime between sets — perfect dwell window."),
        ("STADIUMS",         "Concourse activations.",        "Fan zones, pre-game lines, halftime traffic."),
        ("RETAIL & MALLS",   "Floor + window installs.",      "Walk-by traffic, especially in foodcourts and atriums."),
        ("STREET TAKEOVERS", "Guerrilla LED truck drops.",    "Pull up, draw a crowd, drive away. Repeat at the next corner."),
        ("CONFERENCES",      "Booths that beat every booth.", "The activation that has a line. The line is the marketing."),
        ("CAMPUS",           "Universities & schools.",       "Quad takeovers, orientation events, recruiting drives."),
        ("LAUNCH EVENTS",    "Brand launches & premieres.",   "Captures press photos. The screen becomes the press shot."),
        ("HOLIDAY MARKETS",  "Seasonal villages, markets.",   "Slow-walking, holiday-mood crowds. High capture rate."),
    ], cols=4)


def slide_deliverables(doc):
    add_eyebrow(doc, "12 · What you get")
    add_title(doc, "{em}Turnkey{em}. Custom-branded.", size=40)
    add_card_row(doc, [
        ("01 · CUSTOM BUILD",
         "Fully reskinned game",
         "Your palette, fonts, logo, hero art, audio, copy. Brief-to-live in days, not months."),
        ("02 · HARDWARE-READY",
         "Any screen, any size",
         "Runs on iPad in landscape, any LED panel dimension, projection-ready output. Hand you a URL — it runs anywhere."),
        ("03 · OPERATOR PANEL",
         "On-site dial-in",
         "Difficulty, throw power, target speed, reaction copy — the activation tunes to the crowd in real time."),
        ("04 · CAPTURE FLOW",
         "Wired to your CRM",
         "QR drop wired into your CRM, sweepstakes platform, email tool, or sample-fulfillment vendor of choice."),
        ("05 · MULTI-STOP REUSE",
         "Built once, runs everywhere",
         "Built once, runs at every stop on a tour. Reset the leaderboard between events. Same URL, same brand."),
        ("06 · REPORTING",
         "Lightweight dashboard",
         "Plays, completions, top score, dwell time, capture conversion. Enough to brief the next stop."),
    ], cols=3)


def slide_demos(doc):
    add_eyebrow(doc, "13 · See it live")
    add_title(doc, "Try it. {em}Right now.{em}", size=48)
    add_lede(doc,
             "Visit interactivemedia.ca on any phone, tablet, or laptop. "
             "Enter one of the codes below to open the live demo.")
    add_card_row(doc, [
        ("ENTER CODE", "DEMO",
         "Side-by-side preview — iPad + truck syncing live in one page."),
        ("ENTER CODE", "DEMO 2",
         "Tap-to-the-beat reflex sprint — a single-surface example."),
        ("ENTER CODE", "OREO",
         "Sample branded build — Sling The Stuf. Shows what a reskin looks like."),
        ("FOR PITCHES", "CUSTOM CODE",
         "We'll issue a private demo code for any active opportunity — branded to the prospect."),
    ], cols=4)
    add_spacer(doc, 6)
    p = doc.add_paragraph()
    add_run(p, "Each demo runs in any modern browser — no app, no download, no login.",
            font="Courier New", size=9, color=INK3, letter_spacing=60)


def slide_cta(doc):
    add_eyebrow(doc, "Next steps")
    add_title(doc, "Pitch us a brand. {em}We'll send a custom demo back.{em}", size=32)
    add_lede(doc,
             "Bring us a brief — even a rough one — and we'll return a "
             "live, branded build inside a working week. Use it to close the "
             "client. Take it on the road from there.")
    add_section_label(doc, "Contact — fill in before sending")
    add_bullet(doc, "Rep",   "[Your name]")
    add_bullet(doc, "Email", "[your.email@interactivemedia.ca]")
    add_bullet(doc, "Phone", "[+1 · your phone]")
    add_bullet(doc, "Web",   "interactivemedia.ca")
    add_bullet(doc, "QR",    "Drop a vCard QR into this corner before printing.")


# ──────────────────────────────────────────────────────────────────────────────
# Compose
# ──────────────────────────────────────────────────────────────────────────────
SLIDES = [
    ("Cover",                "",                            slide_cover),
    ("The problem",          "01 · PROBLEM",                slide_problem),
    ("The product",          "02 · PRODUCT",                slide_idea),
    ("Why it works",         "03 · WHY",                    slide_why),
    ("The setup",            "04 · SETUP",                  slide_setup),
    ("The iPad",             "05 · SURFACE — PLAYER",       slide_ipad),
    ("The side LED",         "06 · SURFACE — HYPE",         slide_side_led),
    ("The rear panel",       "07 · SURFACE — CAPTURE",      slide_rear),
    ("Game ideas (action)",  "08 · IDEAS",                  slide_game_ideas_action),
    ("Game ideas (brand)",   "08 · IDEAS CONT.",            slide_game_ideas_brand),
    ("Worked example",       "09 · EXAMPLE",                slide_example_walkthrough),
    ("Brand reskins",        "10 · RESKINS",                slide_examples_brands),
    ("Where it goes",        "11 · CONTEXTS",               slide_contexts),
    ("What you get",         "12 · DELIVERABLES",           slide_deliverables),
    ("Demo codes",           "13 · DEMOS",                  slide_demos),
    ("CTA",                  "NEXT STEPS",                  slide_cta),
]


def build_docx(path: Path):
    doc = Document()
    landscape_section(doc)
    set_page_bg(doc, BG)

    # remove the empty default paragraph at the very top so the first
    # header table is the first thing on page 1 (if it exists)
    if doc.paragraphs:
        first = doc.paragraphs[0]
        first._element.getparent().remove(first._element)

    total = len(SLIDES)
    for i, (_title, section, fn) in enumerate(SLIDES, start=1):
        slide_header(doc, i, total, section)
        fn(doc)
        if i < total:
            page_break(doc)

    doc.save(str(path))


if __name__ == "__main__":
    here = Path(__file__).parent
    out = here / "Interactive-Media-Pitch.docx"
    build_docx(out)
    print(f"Wrote {out}")
